#!/usr/bin/env python3
"""
Generate a single Word document that combines the functional design (FDD) and
the technical architecture for meet.witysk.org. Embeds matplotlib-rendered
diagrams (system architecture, ER, sequence flows) and schematic UI wireframes.

Run inside the project's translate venv (it has matplotlib + python-docx):

    /tmp/translate_venv/bin/python scripts/build_design_doc.py

Output: docs/meet-fdd-architecture.docx
"""
from __future__ import annotations

import io
import os
from pathlib import Path

import matplotlib.patches as patches
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch

REPO = Path(__file__).resolve().parents[1]
OUT_DIR = REPO / "docs"
OUT_PATH = OUT_DIR / "meet-fdd-architecture.docx"
IMG_DIR = OUT_DIR / "_assets"

PRIMARY = "#1f3a8a"   # navy
ACCENT = "#16a34a"    # green
GREY = "#475569"
LIGHT = "#e2e8f0"
EDGE = "#0f172a"

plt.rcParams.update({
    "font.family": "DejaVu Sans",
    "font.size": 9,
    "axes.spines.top": False,
    "axes.spines.right": False,
    "axes.spines.left": False,
    "axes.spines.bottom": False,
})


# ─────────────────────────────────────────────────────────────────────────────
# Diagram helpers
# ─────────────────────────────────────────────────────────────────────────────
def _save(fig, name: str) -> Path:
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    p = IMG_DIR / f"{name}.png"
    fig.savefig(p, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return p


def _box(ax, x, y, w, h, text, *, fill=LIGHT, edge=EDGE, text_color=EDGE, fs=9, bold=False):
    box = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.04,rounding_size=0.08",
                         linewidth=1.2, edgecolor=edge, facecolor=fill)
    ax.add_patch(box)
    weight = "bold" if bold else "normal"
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center",
            fontsize=fs, color=text_color, fontweight=weight, wrap=True)


def _arrow(ax, x1, y1, x2, y2, label=None, *, color=EDGE, ls="-", lw=1.0, label_offset=(0.0, 0.10)):
    arr = FancyArrowPatch((x1, y1), (x2, y2),
                          arrowstyle="->", mutation_scale=10,
                          linewidth=lw, color=color, linestyle=ls,
                          connectionstyle="arc3,rad=0.0")
    ax.add_patch(arr)
    if label:
        ax.text((x1 + x2) / 2 + label_offset[0], (y1 + y2) / 2 + label_offset[1],
                label, ha="center", va="center", fontsize=7, color=color,
                bbox=dict(facecolor="white", edgecolor="none", pad=1.0))


def render_architecture():
    fig, ax = plt.subplots(figsize=(11, 7))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9)
    ax.set_aspect("equal")
    ax.axis("off")

    ax.text(7, 8.6, "meet.witysk.org — system architecture",
            ha="center", fontsize=12, fontweight="bold")

    # Browser
    _box(ax, 0.2, 6.0, 2.6, 1.4, "Browser SPA\nReact + Vite\n@livekit/components", fill="#dbeafe", bold=True)
    # one.witysk.org SSO
    _box(ax, 0.2, 3.8, 2.6, 1.0, "one.witysk.org\nSSO bootstrap iframe", fill="#fde68a")
    # coturn
    _box(ax, 0.2, 1.6, 2.6, 1.0, "coturn (TURN/STUN)\nturn.witysk.org\n(separate process)", fill="#fee2e2")

    # Caddy edge
    _box(ax, 4.0, 6.0, 2.6, 1.4, "Caddy 2\nTLS + reverse proxy\n443/80/UDP-443", fill="#bbf7d0", bold=True)

    # Internal services column
    _box(ax, 7.6, 7.0, 2.6, 1.0, "frontend-build\n(static SPA bundle)", fill=LIGHT)
    _box(ax, 7.6, 5.6, 2.6, 1.0, "meeting-api\nFastAPI + SQLAlchemy", fill="#c7d2fe", bold=True)
    _box(ax, 7.6, 4.2, 2.6, 1.0, "LiveKit server\nWebRTC SFU", fill="#c7d2fe")
    _box(ax, 7.6, 2.8, 2.6, 1.0, "LiveKit Egress\nrecording worker", fill="#c7d2fe")
    _box(ax, 7.6, 1.4, 2.6, 1.0, "Redis 7\n(LiveKit state)", fill=LIGHT)

    # Storage / external column
    _box(ax, 11.2, 7.0, 2.6, 1.0, "SQLite\n/var/lib/meet/meet.db", fill="#fef3c7")
    _box(ax, 11.2, 5.6, 2.6, 1.0, "/var/lib/meet/recordings/\n(local files, 30d TTL)", fill="#fef3c7")
    _box(ax, 11.2, 4.2, 2.6, 1.0, "YouTube Data API\n(owner-initiated upload)", fill="#fee2e2")
    _box(ax, 11.2, 2.8, 2.6, 1.0, "Resend\n(invite emails)", fill="#fee2e2")
    _box(ax, 11.2, 1.4, 2.6, 1.0, "DeepSeek API\n(translation tooling, dev)", fill="#fee2e2")

    # Arrows
    _arrow(ax, 2.8, 6.7, 4.0, 6.7, "HTTPS / WSS", lw=1.4)
    _arrow(ax, 2.8, 4.3, 4.0, 6.5, "postMessage\n(access_token)", color=GREY, ls="--")
    _arrow(ax, 2.8, 2.1, 7.6, 4.4, "ICE / TURN\n(UDP, fallback TCP/TLS)", color=GREY, ls="--")

    _arrow(ax, 6.6, 7.0, 7.6, 7.4, "/ → static")
    _arrow(ax, 6.6, 6.7, 7.6, 6.0, "/api/*")
    _arrow(ax, 6.6, 6.3, 7.6, 4.6, "/rtc/* (WSS)")

    _arrow(ax, 10.2, 6.0, 11.2, 7.3, "ORM")
    _arrow(ax, 10.2, 5.8, 11.2, 5.9, "rec/* download")
    _arrow(ax, 10.2, 5.6, 11.2, 4.5, "publish (oauth)")
    _arrow(ax, 10.2, 5.4, 11.2, 3.1, "send invite")

    _arrow(ax, 10.2, 4.3, 10.2, 3.6, "egress RPC", color=GREY)
    _arrow(ax, 10.2, 3.8, 10.2, 2.4, "state", color=GREY)
    _arrow(ax, 10.2, 4.4, 10.2, 1.9, "presence", color=GREY, ls=":")

    return _save(fig, "architecture")


def render_er():
    fig, ax = plt.subplots(figsize=(11, 6.5))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9)
    ax.set_aspect("equal")
    ax.axis("off")
    ax.text(7, 8.6, "Database schema (SQLite, SQLAlchemy)", ha="center", fontsize=12, fontweight="bold")

    def entity(x, y, w, h, name, fields):
        # header
        head = FancyBboxPatch((x, y + h - 0.55), w, 0.55,
                              boxstyle="round,pad=0.02,rounding_size=0.05",
                              linewidth=1.2, edgecolor=EDGE, facecolor=PRIMARY)
        ax.add_patch(head)
        ax.text(x + w / 2, y + h - 0.28, name, ha="center", va="center",
                color="white", fontsize=10, fontweight="bold")
        # body
        body = patches.Rectangle((x, y), w, h - 0.55, linewidth=1.2,
                                 edgecolor=EDGE, facecolor="white")
        ax.add_patch(body)
        for i, f in enumerate(fields):
            ax.text(x + 0.10, y + h - 0.55 - 0.20 - i * 0.22, f,
                    ha="left", va="center", fontsize=7.5, color=EDGE)

    entity(0.3, 4.4, 4.5, 4.0, "Meeting",
           ["id  PK  (ULID)",
            "room_name  UNIQUE  (3-word slug)",
            "display_title",
            "owner_user_id  IDX  (sub from JWT)",
            "owner_email / owner_name",
            "is_active  /  closed_at  /  hidden",
            "require_password  /  password_hash",
            "max_participants  (2..50)",
            "recording_mode  (manual|auto|off)",
            "branding_image_path",
            "list_for_authenticated / list_for_anonymous",
            "created_at / scheduled_at / ends_at"])

    entity(5.2, 6.5, 4.0, 1.9, "MeetingParticipant",
           ["id PK", "meeting_id FK → Meeting",
            "livekit_identity  IDX",
            "display_name / email",
            "is_authenticated / is_owner",
            "joined_at / left_at"])

    entity(5.2, 4.0, 4.0, 2.2, "Recording",
           ["id PK  (ULID)", "meeting_id FK", "egress_id  IDX",
            "file_path / file_size_bytes / duration_seconds",
            "started_at / ended_at / expires_at  IDX",
            "status (running|completed|failed|deleted)",
            "youtube_url / youtube_video_id",
            "youtube_status / youtube_error"])

    entity(5.2, 1.6, 4.0, 2.0, "ModerationAudit",
           ["id PK", "meeting_id FK",
            "actor_user_id", "action (mute|kick|presenter|…)",
            "target_identity", "details", "created_at"])

    entity(9.7, 4.4, 4.0, 2.0, "UserPreferences",
           ["user_id  PK  (sub from JWT)",
            "language  (e.g. 'de', nullable)",
            "language_set_manually  BOOL",
            "created_at / updated_at"])

    # FK arrows from children → Meeting
    _arrow(ax, 5.2, 7.3, 4.8, 7.5, "FK")
    _arrow(ax, 5.2, 5.0, 4.8, 6.0, "FK")
    _arrow(ax, 5.2, 2.5, 4.8, 5.0, "FK")

    # Note about UserPreferences (no FK to Meeting; keyed by JWT sub)
    ax.text(11.7, 3.9, "(no FK; keyed by\nJWT sub claim)", ha="center", fontsize=7, style="italic", color=GREY)

    return _save(fig, "er-diagram")


def _seq(ax, x):
    """Draw a simple lifeline at column x."""
    ax.plot([x, x], [0.5, 8.5], ls="--", color=GREY, lw=0.8)


def render_sso_sequence():
    fig, ax = plt.subplots(figsize=(11, 7))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 10)
    ax.axis("off")
    ax.text(7, 9.4, "SSO bootstrap & API auth (one.witysk.org → meet.witysk.org)",
            ha="center", fontsize=11, fontweight="bold")

    cols = [("Browser\n(SPA)", 1.5),
            ("hidden iframe\nsso-bootstrap.html", 4.5),
            ("one.witysk.org\nlocalStorage", 7.5),
            ("meet-api\nFastAPI", 10.5),
            ("Caddy\nproxy", 12.8)]
    for label, x in cols:
        _box(ax, x - 1.1, 8.5, 2.2, 0.6, label, fill=PRIMARY, text_color="white", fs=8, bold=True)
        ax.plot([x, x], [0.3, 8.5], ls="--", color=GREY, lw=0.8)

    def step(y, x1, x2, label, *, ls="-", color=EDGE):
        ax.annotate("", xy=(x2, y), xytext=(x1, y),
                    arrowprops=dict(arrowstyle="->", color=color, linestyle=ls, lw=1.0))
        ax.text((x1 + x2) / 2, y + 0.16, label, ha="center", fontsize=7.5, color=color)

    # Steps top-down
    step(7.8, 1.5, 4.5, "1. App.useEffect: bootstrapFromOneWitysk()")
    step(7.2, 4.5, 7.5, "2. iframe loads sso-bootstrap.html (cross-origin)", color=GREY)
    step(6.6, 7.5, 4.5, "3. read localStorage.access_token  → postMessage(access_token)", color=GREY)
    step(6.0, 4.5, 1.5, "4. message event → mirror to meet localStorage")
    step(5.4, 1.5, 1.5, "5. syncServerLanguage(): pick up token", ls=":")
    step(4.8, 1.5, 12.8, "6. fetch('/api/v1/me/preferences', Authorization: Bearer)")
    step(4.2, 12.8, 10.5, "7. proxy → meeting-api")
    step(3.6, 10.5, 10.5, "8. RequireUser → decode HS256(JWT)")
    step(3.0, 10.5, 1.5, "9. {language, language_set_manually}", color=ACCENT)
    step(2.4, 1.5, 1.5, "10. if manual ⇒ i18n.changeLanguage(server_lang)", ls=":", color=ACCENT)

    return _save(fig, "seq-sso")


def render_language_sequence():
    fig, ax = plt.subplots(figsize=(11, 6.5))
    ax.set_xlim(0, 14); ax.set_ylim(0, 10); ax.axis("off")
    ax.text(7, 9.4, "Language preference: pick & persist",
            ha="center", fontsize=11, fontweight="bold")

    cols = [("User", 1.0), ("Settings page\n(pref-lang)", 3.5),
            ("preferences\nstore (zustand)", 6.5), ("i18next", 9.0),
            ("meet-api\n/me/preferences", 12.0)]
    for label, x in cols:
        _box(ax, x - 1.0, 8.5, 2.0, 0.6, label, fill=PRIMARY, text_color="white", fs=8, bold=True)
        ax.plot([x, x], [0.3, 8.5], ls="--", color=GREY, lw=0.8)

    def step(y, x1, x2, label, *, ls="-", color=EDGE):
        ax.annotate("", xy=(x2, y), xytext=(x1, y),
                    arrowprops=dict(arrowstyle="->", color=color, linestyle=ls, lw=1.0))
        ax.text((x1 + x2) / 2, y + 0.16, label, ha="center", fontsize=7.5, color=color)

    step(7.8, 1.0, 3.5, "1. select 'Deutsch'")
    step(7.2, 3.5, 6.5, "2. setLocale({language: 'de'})")
    step(6.6, 6.5, 9.0, "3. notifyLanguageChange → i18n.changeLanguage('de')")
    step(6.0, 9.0, 3.5, "4. languageChanged → React re-renders (t() pulls de.json)", color=ACCENT)
    step(5.2, 6.5, 12.0, "5. PUT /api/v1/me/preferences {language: 'de'}", ls=":")
    step(4.6, 12.0, 6.5, "6. {language: 'de', language_set_manually: TRUE}", color=ACCENT)
    step(3.6, 1.0, 1.0, "—— next visit ——", ls=":", color=GREY)
    step(3.0, 1.0, 12.0, "7. App boot: GET /me/preferences (after SSO bootstrap)")
    step(2.4, 12.0, 9.0, "8. server returns {de, manual=true} → i18n.changeLanguage('de')", color=ACCENT)
    step(1.6, 1.0, 1.0, "If language_set_manually=false ⇒ keep browser-detected lang", ls=":", color=GREY)

    return _save(fig, "seq-language")


def render_recording_sequence():
    fig, ax = plt.subplots(figsize=(11, 7))
    ax.set_xlim(0, 14); ax.set_ylim(0, 10); ax.axis("off")
    ax.text(7, 9.4, "Recording lifecycle (start → file → publish → expire)",
            ha="center", fontsize=11, fontweight="bold")

    cols = [("Owner\n(SPA)", 1.0), ("meeting-api", 4.0), ("LiveKit\nserver", 7.0),
            ("Egress\nworker", 9.5), ("Local FS\n+ DB", 12.5)]
    for label, x in cols:
        _box(ax, x - 1.1, 8.5, 2.2, 0.6, label, fill=PRIMARY, text_color="white", fs=8, bold=True)
        ax.plot([x, x], [0.3, 8.5], ls="--", color=GREY, lw=0.8)

    def step(y, x1, x2, label, *, ls="-", color=EDGE):
        ax.annotate("", xy=(x2, y), xytext=(x1, y),
                    arrowprops=dict(arrowstyle="->", color=color, linestyle=ls, lw=1.0))
        ax.text((x1 + x2) / 2, y + 0.16, label, ha="center", fontsize=7.5, color=color)

    step(7.9, 1.0, 4.0, "1. POST /meetings/{id}/recordings:start")
    step(7.3, 4.0, 7.0, "2. set room metadata.recording_active=true")
    step(6.7, 4.0, 9.5, "3. egress.start_room_composite(...)")
    step(6.1, 9.5, 12.5, "4. write segments, mux mp4")
    step(5.5, 9.5, 4.0, "5. webhook: egress_ended", color=ACCENT)
    step(4.9, 4.0, 12.5, "6. INSERT recordings (status=completed, file_path)", color=ACCENT)
    step(4.0, 1.0, 4.0, "—— later ——", ls=":", color=GREY)
    step(3.4, 1.0, 4.0, "7. Click Publish → POST /recordings/{id}/publish-youtube")
    step(2.8, 4.0, 12.5, "8. read file → upload to YouTube → set youtube_url, delete local")
    step(2.0, 1.0, 12.5, "9. scheduler sweeps expires_at < now → mark deleted, rm file", ls=":", color=GREY)

    return _save(fig, "seq-recording")


# ─────────────────────────────────────────────────────────────────────────────
# Wireframes (schematic, not pixel-perfect)
# ─────────────────────────────────────────────────────────────────────────────
def _wf_setup(title, w=12, h=7):
    fig, ax = plt.subplots(figsize=(w, h))
    ax.set_xlim(0, 12); ax.set_ylim(0, 8); ax.set_aspect("equal"); ax.axis("off")
    ax.text(6, 7.7, title, ha="center", fontsize=11, fontweight="bold")
    return fig, ax


def _wf_sidebar(ax, x=0.0):
    _box(ax, x, 0.4, 1.6, 7.0, "", fill="#1e293b", edge="#1e293b")
    ax.text(x + 0.8, 7.0, "meet", ha="center", color="white", fontweight="bold")
    ax.text(x + 0.8, 6.78, "witysk.org", ha="center", color="#cbd5e1", fontsize=7)
    items = [("● Home", 6.2, True), ("○ Recordings", 5.7, False), ("○ Settings", 5.2, False)]
    for label, y, sel in items:
        if sel:
            _box(ax, x + 0.1, y - 0.15, 1.4, 0.35, "", fill="#334155", edge="#334155")
        ax.text(x + 0.8, y, label, ha="center", color="white", fontsize=8.5)
    # footer
    ax.text(x + 0.8, 1.3, "✓ signed in", ha="center", color="#22c55e", fontsize=7)
    ax.text(x + 0.8, 1.05, "via one.witysk.org", ha="center", color="#94a3b8", fontsize=6.5)
    _box(ax, x + 0.2, 0.5, 1.2, 0.30, "Log off", fill="#1e293b", edge="#475569", text_color="white", fs=7)


def render_wf_home():
    fig, ax = _wf_setup("Home / Create meeting (CreateMeeting.tsx + MyMeetings + Discover)")
    _wf_sidebar(ax)

    # My meetings card
    _box(ax, 1.9, 5.3, 9.9, 2.0, "", fill="white", edge=GREY)
    ax.text(2.1, 7.05, "Active meetings", fontsize=9, fontweight="bold")
    for i, (title, slug) in enumerate([("Weekly sync", "happy-blue-tiger"),
                                       ("1:1 with Alice", "noisy-quiet-sun")]):
        y = 6.6 - i * 0.55
        _box(ax, 2.1, y - 0.20, 4.0, 0.40, f"CAM  {title}", fill=LIGHT, fs=8)
        ax.text(6.2, y, slug, color=GREY, fontsize=7)
        _box(ax, 8.2, y - 0.20, 0.5, 0.40, "Cp", fs=8)
        _box(ax, 8.8, y - 0.20, 0.5, 0.40, "@", fs=10)
        _box(ax, 9.4, y - 0.20, 0.85, 0.40, "Join", fill=ACCENT, text_color="white", fs=8, bold=True)
        _box(ax, 10.4, y - 0.20, 0.7, 0.40, "End", fill="#dc2626", text_color="white", fs=8)

    # Closed meetings (with Restart button — new)
    _box(ax, 1.9, 3.6, 9.9, 1.5, "", fill="white", edge=GREY)
    ax.text(2.1, 4.9, "Closed meetings", fontsize=9, fontweight="bold", color=GREY)
    _box(ax, 2.1, 4.30, 4.0, 0.30, "Q3 review", fill=LIGHT, fs=8)
    ax.text(6.2, 4.45, "merry-orange-cat", color=GREY, fontsize=7)
    ax.text(8.2, 4.45, "closed", color=GREY, fontsize=7)
    _box(ax, 8.8, 4.27, 1.4, 0.36, "Restart", fill=ACCENT, text_color="white", fs=8, bold=True)
    _box(ax, 10.3, 4.27, 0.6, 0.36, "Del", fs=8)

    # Discover
    _box(ax, 1.9, 2.0, 9.9, 1.4, "", fill="white", edge=GREY)
    ax.text(2.1, 3.20, "Discover — meetings open to join", fontsize=9, fontweight="bold")
    _box(ax, 2.1, 2.55, 4.0, 0.40, "Open standup", fill=LIGHT, fs=8)
    ax.text(6.2, 2.75, "puzzle-radio-leaf", color=GREY, fontsize=7)
    _box(ax, 9.4, 2.55, 0.85, 0.40, "Join", fill=ACCENT, text_color="white", fs=8, bold=True)

    # Create form
    _box(ax, 1.9, 0.4, 9.9, 1.3, "", fill="white", edge=GREY)
    ax.text(2.1, 1.45, "Create a meeting", fontsize=9, fontweight="bold")
    _box(ax, 2.1, 0.85, 6.0, 0.35, "Title …", fill=LIGHT, fs=8)
    _box(ax, 2.1, 0.5, 1.4, 0.30, "☐ Password", fs=7)
    _box(ax, 9.0, 0.7, 2.1, 0.50, "Create meeting", fill=PRIMARY, text_color="white", fs=8, bold=True)

    return _save(fig, "wf-home")


def render_wf_lobby():
    fig, ax = _wf_setup("Lobby / join page (Lobby.tsx)")
    _wf_sidebar(ax)
    # central card
    _box(ax, 3.5, 1.6, 6.0, 5.4, "", fill="white", edge=GREY)
    _box(ax, 3.7, 5.6, 1.0, 1.0, "img", fill=LIGHT, fs=8)
    ax.text(4.9, 6.55, "Weekly sync", fontsize=11, fontweight="bold")
    ax.text(4.9, 6.25, "Room: happy-blue-tiger", color=GREY, fontsize=8)
    ax.text(4.9, 5.95, "Hosted by Stéphane", color=GREY, fontsize=8)

    for i, label in enumerate(["Your name", "Email (optional)", "Password (if required)"]):
        y = 4.7 - i * 0.85
        ax.text(3.7, y + 0.30, label, fontsize=8)
        _box(ax, 3.7, y - 0.05, 5.6, 0.35, "", fill=LIGHT, fs=8)

    _box(ax, 3.7, 2.0, 1.6, 0.55, "Join", fill=ACCENT, text_color="white", fs=10, bold=True)
    return _save(fig, "wf-lobby")


def render_wf_room():
    fig, ax = _wf_setup("Room — in-meeting (Room.tsx)")
    # full-width topbar
    _box(ax, 0.0, 6.5, 12.0, 0.9, "", fill="#0f172a", edge="#0f172a")
    ax.text(0.3, 7.10, "CAM  meet.witysk.org   • REC", color="white", fontsize=8.5, fontweight="bold")
    for i, l in enumerate(["BG v", "Invite", "Take stage", "Grid", "Mute all", "Stop rec", "End", "Set", "People", "Chat"]):
        col = "#dc2626" if l in ("Stop rec", "End") else "#16a34a" if l == "Invite" else "#334155"
        _box(ax, 2.7 + i * 0.93, 6.66, 0.85, 0.55, l, fill=col, text_color="white", fs=7)

    # stage
    _box(ax, 0.0, 1.0, 8.6, 5.4, "", fill="#0b1220", edge="#0b1220")
    for r in range(2):
        for c in range(3):
            x = 0.2 + c * 2.78
            y = 4.0 - r * 2.5
            _box(ax, x, y, 2.6, 2.0, "video tile", fill="#1e293b", edge="#475569",
                 text_color="white", fs=8)

    # right panel — chat
    _box(ax, 8.7, 1.0, 3.3, 5.4, "", fill="white", edge=GREY)
    ax.text(8.85, 6.20, "Chat", fontsize=9, fontweight="bold")
    for i in range(3):
        y = 5.4 - i * 0.75
        _box(ax, 8.85, y - 0.30, 3.0, 0.5, "Alice: hi", fill=LIGHT, fs=7)
    _box(ax, 8.85, 1.2, 2.5, 0.40, "Type a message", fill=LIGHT, fs=7)
    _box(ax, 11.40, 1.2, 0.45, 0.40, "→", fill=ACCENT, text_color="white", fs=10, bold=True)

    # bottom control bar
    _box(ax, 0.0, 0.0, 12.0, 0.9, "", fill="#0f172a", edge="#0f172a")
    for i, l in enumerate(["Mic", "Cam", "Share", "Hand", "Leave"]):
        _box(ax, 4.4 + i * 0.85, 0.2, 0.78, 0.55, l, fill="#1e293b", text_color="white", fs=7.5)
    return _save(fig, "wf-room")


def render_wf_settings():
    fig, ax = _wf_setup("Settings — Language tab (Settings.tsx, locale)")
    _wf_sidebar(ax)

    ax.text(2.0, 7.20, "Settings", fontsize=12, fontweight="bold")
    ax.text(2.0, 6.92, "Preferences are stored in this browser. Auth users sync language to the server.",
            fontsize=7.5, color=GREY)

    # tab strip
    tabs = ["Audio & Video", "Display", "Meeting", "Moderation", "Recording",
            "Notifications", "Privacy", "Accessibility", "Keyboard", "Network",
            "Language", "Chat", "Appearance", "Developer", "Reset"]
    x = 2.0; y = 6.3
    for i, t in enumerate(tabs):
        sel = (t == "Language")
        col = PRIMARY if sel else "#334155"
        tx = 2.0 + (i % 8) * 1.18
        ty = 6.30 - (i // 8) * 0.45
        _box(ax, tx, ty, 1.05, 0.32, t, fill=col, text_color="white", fs=6.5, bold=sel)

    # form card
    _box(ax, 2.0, 0.8, 9.5, 4.6, "", fill="white", edge=GREY)
    ax.text(2.2, 5.10, "Language & locale", fontsize=10, fontweight="bold")

    ax.text(2.2, 4.65, "Interface language", fontsize=8)
    _box(ax, 2.2, 4.10, 6.5, 0.40, "Deutsch (German)         ▼", fill=LIGHT, fs=8.5)

    ax.text(2.2, 3.65, "Time format", fontsize=8)
    _box(ax, 2.2, 3.20, 4.0, 0.35, "24-hour (13:30)         ▼", fill=LIGHT, fs=8)

    ax.text(2.2, 2.85, "Date format", fontsize=8)
    _box(ax, 2.2, 2.40, 4.0, 0.35, "DD/MM/YYYY               ▼", fill=LIGHT, fs=8)

    ax.text(2.2, 2.10, "Timezone", fontsize=8)
    _box(ax, 2.2, 1.65, 4.0, 0.35, "Europe/Brussels", fill=LIGHT, fs=8)

    # Note about server sync
    _box(ax, 7.0, 2.20, 4.3, 1.50, "", fill="#dcfce7", edge="#16a34a")
    ax.text(9.15, 3.45, "(i)  Server sync (auth users)", fontsize=8, fontweight="bold", ha="center", color="#166534")
    ax.text(9.15, 3.10, "PUT /me/preferences", fontsize=7, ha="center", color="#166534")
    ax.text(9.15, 2.80, "language_set_manually = TRUE", fontsize=7, ha="center", color="#166534")
    ax.text(9.15, 2.50, "Browser-language detection", fontsize=7, ha="center", color="#166534")
    ax.text(9.15, 2.30, "is suppressed on next login.", fontsize=7, ha="center", color="#166534")

    return _save(fig, "wf-settings")


def render_wf_recordings():
    fig, ax = _wf_setup("Recordings page (Recordings.tsx)")
    _wf_sidebar(ax)

    ax.text(2.0, 7.2, "Recordings", fontsize=12, fontweight="bold")
    ax.text(2.0, 6.92, "Server-side recordings. Local files auto-delete after 30 days.",
            fontsize=7.5, color=GREY)

    for i, (rid, title, status, yt) in enumerate([
        ("01J...A2", "Weekly sync — 2026-04-23", "completed", None),
        ("01J...B7", "All-hands", "completed", "published"),
        ("01J...C9", "Customer call", "completed", "uploading"),
    ]):
        y = 5.5 - i * 1.4
        _box(ax, 2.0, y - 0.4, 9.5, 1.2, "", fill="white", edge=GREY)
        _box(ax, 2.2, y, 0.7, 0.7, "CAM", fill=LIGHT, fs=10)
        ax.text(3.0, y + 0.45, rid, fontsize=8, color=GREY)
        ax.text(4.6, y + 0.45, title, fontsize=9, fontweight="bold")
        _box(ax, 4.6, y + 0.05, 0.95, 0.32, status, fill="#dcfce7", text_color="#166534", fs=7)
        if yt:
            colour = "#fde68a" if yt == "uploading" else "#fecaca"
            _box(ax, 5.65, y + 0.05, 1.15, 0.32, f"YT: {yt}", fill=colour, fs=7)
        ax.text(3.0, y + 0.10, "Started 2026-04-23 14:30 · 47m · 612 MB", fontsize=7, color=GREY)
        # actions
        _box(ax, 7.4, y + 0.05, 1.1, 0.5, "Download", fill="#334155", text_color="white", fs=7)
        _box(ax, 8.6, y + 0.05, 1.6, 0.5, "Publish to YouTube", fill=ACCENT, text_color="white", fs=7, bold=True)
        _box(ax, 10.3, y + 0.05, 0.6, 0.5, "Del", fs=8)

    return _save(fig, "wf-recordings")


# ─────────────────────────────────────────────────────────────────────────────
# Document assembly
# ─────────────────────────────────────────────────────────────────────────────
def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, *, italic=False, bold=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    if size:
        r.font.size = Pt(size)
    return p


def add_bullets(doc, items):
    for i in items:
        doc.add_paragraph(i, style="List Bullet")


def add_image(doc, path, *, width_cm=16.5, caption=None):
    doc.add_picture(str(path), width=Cm(width_cm))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x47, 0x55, 0x69)


def add_table(doc, headers, rows, *, col_widths_cm=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = str(val)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if col_widths_cm:
        for row in t.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)
    return t


def add_toc_field(doc):
    """Insert a Word TOC field. The user must press F9 (or right-click → Update field)
    once to populate it. We place a placeholder paragraph below."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:t")
    fldChar3.text = "Right-click → Update Field to populate the table of contents."
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    run._r.append(fldChar4)


def build_doc():
    print("→ rendering diagrams …")
    arch_png = render_architecture()
    er_png = render_er()
    sso_png = render_sso_sequence()
    lang_png = render_language_sequence()
    rec_png = render_recording_sequence()

    print("→ rendering wireframes …")
    wf_home = render_wf_home()
    wf_lobby = render_wf_lobby()
    wf_room = render_wf_room()
    wf_settings = render_wf_settings()
    wf_recordings = render_wf_recordings()

    print("→ assembling docx …")
    doc = Document()

    # Default style tweak — slightly tighter
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ── Cover ────────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("meet.witysk.org")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x1f, 0x3a, 0x8a)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Functional Design & Technical Architecture")
    r.font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("Self-hosted video conferencing — internal design document")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    doc.add_paragraph()
    doc.add_paragraph()
    add_table(doc,
              ["Item", "Value"],
              [["Document version", "1.0"],
               ["Status", "Living document"],
               ["Owner", "Stéphane (stepvda)"],
               ["Production host", "turn.witysk.org (Caddy + Docker compose project 'meet')"],
               ["Public URL", "https://meet.witysk.org"],
               ["Source repo", "github local — /opt/meet on the deploy host"]],
              col_widths_cm=[5.0, 11.5])
    doc.add_page_break()

    # ── TOC ──────────────────────────────────────────────────────────────────
    add_heading(doc, "Table of contents", level=1)
    add_toc_field(doc)
    doc.add_page_break()

    # ── 1. Introduction ──────────────────────────────────────────────────────
    add_heading(doc, "1. Introduction", level=1)
    add_heading(doc, "1.1 Purpose", level=2)
    add_para(doc,
        "This document describes both the user-facing functionality and the "
        "technical architecture of meet.witysk.org. It is intended as the "
        "single reference an engineer or product owner needs to (a) understand "
        "what the system does, (b) reason about how to extend it, and (c) "
        "deploy and operate it safely.")

    add_heading(doc, "1.2 Scope", level=2)
    add_para(doc,
        "Covered: the SPA, the meeting-api control plane, the LiveKit media "
        "stack, persistent storage, SSO integration with one.witysk.org, the "
        "internationalization design, the recording lifecycle, and the "
        "deployment topology on a single Docker host.")
    add_para(doc,
        "Out of scope: the implementation of one.witysk.org itself, internal "
        "details of LiveKit's WebRTC stack, and any cross-region or "
        "high-availability story (current deployment is a single VM behind "
        "Caddy, alongside coturn).",
        italic=True)

    add_heading(doc, "1.3 Audience", level=2)
    add_bullets(doc, [
        "Engineers maintaining or extending meet.witysk.org.",
        "Product/PM owners reasoning about feature scope.",
        "Operators running the service on the Docker host.",
        "Reviewers performing security or compliance audits.",
    ])

    add_heading(doc, "1.4 Glossary", level=2)
    add_table(doc,
              ["Term", "Definition"],
              [["SPA", "The React/Vite single-page app served from /srv/frontend by Caddy."],
               ["meeting-api", "FastAPI control plane that owns meetings, tokens, recordings, moderation, user prefs."],
               ["LiveKit", "WebRTC SFU. We run livekit-server + livekit-egress as containers."],
               ["coturn", "TURN/STUN relay running on turn.witysk.org. It is intentionally NOT in the meet compose project."],
               ["Owner", "An authenticated user who created (or imported) a meeting."],
               ["Anon / guest", "An unauthenticated participant who joined via a shared link."],
               ["Slug / room_name", "Three-word dash-separated string used as both LiveKit room id and the share URL path."],
               ["one.witysk.org", "External SSO. Issues HS256 JWTs validated by meeting-api."],
               ["Trans wrapper", "i18next <Trans> component that lets translations include inline React elements."]],
              col_widths_cm=[3.5, 13.0])

    doc.add_page_break()

    # ── 2. Functional Design ────────────────────────────────────────────────
    add_heading(doc, "2. Functional Design", level=1)

    add_heading(doc, "2.1 Personas", level=2)
    add_table(doc,
              ["Persona", "Authentication", "Capabilities"],
              [["Owner", "JWT from one.witysk.org",
                "Create meetings; moderate (mute/kick/presenter); start/stop recording; publish to YouTube; restart closed meetings; manage user prefs (incl. language)."],
               ["Authenticated participant", "JWT from one.witysk.org",
                "Discover and join listed meetings; participate in chat, reactions; user prefs sync to backend."],
               ["Anonymous guest", "None — short-lived LiveKit token only",
                "Join a meeting via a shared link; provide a display name; participate (subject to owner moderation)."]],
              col_widths_cm=[3.5, 4.5, 8.5])

    add_heading(doc, "2.2 Feature catalog", level=2)
    add_table(doc,
              ["Area", "Feature", "Notes"],
              [["Meetings", "Create / list / end / delete / reopen", "DELETE has dual semantics: ends if active, soft-hides if closed. POST /reopen flips is_active back, keeping the same room_name (share link)."],
               ["Meetings", "Branding image upload", "JPEG/PNG/WebP/GIF up to 2 MB; served publicly via /api/v1/rooms/{name}/branding."],
               ["Meetings", "Discoverability", "Per-meeting flags: list_for_authenticated, list_for_anonymous (anon implies auth)."],
               ["Meetings", "Hosted-by attribution", "owner_name snapshot fetched from one.witysk.org and shown in lobby + Discover."],
               ["Lobby", "Anonymous join with display name + optional email + optional password", "Mints short-lived LiveKit token."],
               ["Lobby", "Auto-detect ownership", "If JWT user owns meeting, owner token is minted."],
               ["Room", "Audio/video controls, screen share", "Standard LiveKit ControlBar."],
               ["Room", "Background effects", "MediaPipe selfie segmenter via @livekit/track-processors. Built-ins + user upload up to 4 MB."],
               ["Room", "Layouts: grid / spotlight (presenter) / speaker", "Owner sets metadata.presenter_identity to drive spotlight."],
               ["Room", "Recording (owner)", "Server-side via livekit-egress; status mirrored in room metadata."],
               ["Room", "Chat with attachments + reactions + reply", "Persisted server-side; LiveKit data channel signals refetch."],
               ["Room", "Moderation", "Mute (1/all), kick, set/clear presenter, end-for-everyone."],
               ["Recordings", "Download (auth-gated)", "Streamed via /api/v1/recordings/{id}/download with Bearer auth."],
               ["Recordings", "Publish to YouTube", "Owner-initiated; on success the local file is deleted."],
               ["Recordings", "Auto-expiry", "Default 30 days; sweeper job marks deleted and removes files."],
               ["Settings", "15 tabs", "Stored locally; some sync to backend (currently: language)."],
               ["Settings", "Language picker", "50 locales; honors language_set_manually flag (server-side)."],
               ["Settings", "Reset to defaults", "Per-browser only."],
               ["i18n", "49 non-English locales auto-translated via DeepSeek", "Done by scripts/translate_locales.py and scripts/translate_missing_keys.py."]],
              col_widths_cm=[2.5, 4.5, 9.5])

    add_heading(doc, "2.3 User journeys", level=2)

    add_heading(doc, "2.3.1 Owner: SSO, create, record, end, reopen", level=3)
    add_para(doc,
        "An authenticated user lands on meet.witysk.org/. The SPA bootstraps "
        "an access token from one.witysk.org via a hidden iframe, then calls "
        "GET /api/v1/me/preferences. They create a meeting; the API allocates "
        "a unique three-word slug and returns a join URL of the form "
        "meet.witysk.org/<slug>. They join, start a recording, end the "
        "meeting (which marks is_active=False and closes the LiveKit room). "
        "Later, they click Restart on the closed row; the meeting flips back "
        "to active and the original share link continues to work.")

    add_heading(doc, "2.3.2 Anonymous join", level=3)
    add_para(doc,
        "A guest opens a shared meet.witysk.org/<slug> URL. The lobby fetches "
        "public room info (title, branding, host name, password requirement). "
        "They enter their name (and password if required), submit, and are "
        "redirected to /r/<slug> with a short-lived LiveKit token in "
        "sessionStorage. They are not authenticated against meeting-api at "
        "any point; their only interaction with the API is the public lobby "
        "call and the anonymous token mint.")

    add_heading(doc, "2.3.3 Language preference & server persistence", level=3)
    add_para(doc,
        "On first visit, the SPA picks a language using i18next's standard "
        "detector chain: query string → localStorage cache → navigator.language → "
        "fallback English. After SSO bootstrap, App.tsx calls "
        "syncServerLanguage(), which calls GET /me/preferences. If the server "
        "says language_set_manually=true and a language is set, that overrides "
        "the locally-detected one. When the user picks a language in Settings, "
        "the picker writes through to the backend with PUT /me/preferences, "
        "which sets language_set_manually=true. From that point on, the user's "
        "choice follows them across browsers and devices.")
    add_image(doc, lang_png, caption="Figure 2.1 — Language preference flow.")

    add_heading(doc, "2.4 UI wireframes", level=2)
    add_para(doc,
        "Wireframes are schematic — they show the structure and key controls of "
        "each screen rather than final visuals. The production look is dark, "
        "primary-blue with green accents.",
        italic=True)
    add_image(doc, wf_home,        caption="Figure 2.2 — Home page (CreateMeeting + MyMeetings + Discover). Shows the green Restart button on closed rows.")
    add_image(doc, wf_lobby,       caption="Figure 2.3 — Lobby join page. Anon users fill in name + optional email/password.")
    add_image(doc, wf_room,        caption="Figure 2.4 — In-meeting view. Owner-only buttons (Mute all / Stop rec / End) appear conditionally.")
    add_image(doc, wf_settings,    caption="Figure 2.5 — Settings → Language tab. Auth users have their pick synced to the server.")
    add_image(doc, wf_recordings,  caption="Figure 2.6 — Recordings page. Per-row Download / Publish / Delete actions.")

    doc.add_page_break()

    # ── 3. Technical Architecture ───────────────────────────────────────────
    add_heading(doc, "3. Technical Architecture", level=1)

    add_heading(doc, "3.1 System overview", level=2)
    add_para(doc,
        "The deployment runs as a single Docker compose project named 'meet' "
        "on turn.witysk.org. coturn is intentionally a separate process on "
        "the same host (often via systemd) and is NOT part of the compose "
        "project — see Operations §4. The frontend bundle is built inside a "
        "container at deploy time and served by Caddy from a shared volume.")
    add_image(doc, arch_png, caption="Figure 3.1 — System architecture.")

    add_heading(doc, "3.2 Component descriptions", level=2)
    add_table(doc,
              ["Component", "Image / process", "Responsibilities"],
              [["Caddy", "caddy:2",
                "TLS termination (ACME), reverse proxy. Routes /api/* and /rec/* to meeting-api, /rtc/* to LiveKit, otherwise serves the SPA bundle. Sets the CSP."],
               ["frontend-build", "meet-frontend-build (Vite build)",
                "One-shot container that produces dist/ and writes it to the shared frontend volume. The runtime serving is Caddy."],
               ["meeting-api", "meet-meeting-api (FastAPI + uvicorn)",
                "Owns SQLite, mints LiveKit JWTs, handles meetings/moderation/recordings/chat/users routes, listens for LiveKit egress webhooks."],
               ["LiveKit server", "livekit/livekit-server",
                "WebRTC SFU. network_mode: host so it can advertise the public IP for ICE."],
               ["LiveKit Egress", "livekit/egress",
                "Records rooms via Chromium-based composite egress, writes mp4/webm files to a shared volume."],
               ["Redis", "redis:7-alpine",
                "Used by LiveKit for shared state and message bus."],
               ["coturn", "external (systemd or its own container)",
                "TURN relay. Lives on the same VM but outside the meet compose project so deploys never restart it."],
               ["one.witysk.org", "external service",
                "Issues HS256 JWTs. Hosts the sso-bootstrap.html that meet's SPA loads in a hidden iframe to read the access token."]],
              col_widths_cm=[3.0, 4.0, 9.5])

    add_heading(doc, "3.3 Data model", level=2)
    add_image(doc, er_png, caption="Figure 3.2 — Database schema (SQLite). Foreign keys on meeting_id; UserPreferences keyed by JWT sub.")
    add_para(doc,
        "Migrations: there is no Alembic. New tables are auto-created at "
        "startup via Base.metadata.create_all(). Backward-compatible column "
        "additions to existing tables are handled by lightweight_migrate() "
        "in db.py — idempotent ALTER TABLEs guarded by PRAGMA table_info() "
        "checks. New tables (e.g. user_preferences) need no migration code.",
        italic=True)

    add_heading(doc, "3.4 API surface", level=2)
    add_table(doc,
              ["Method", "Path", "Auth", "Purpose"],
              [["GET",    "/api/health",                                   "—",     "Liveness probe."],
               ["POST",   "/api/v1/meetings",                              "owner", "Create a meeting; returns join_url."],
               ["GET",    "/api/v1/meetings",                              "owner", "List my meetings (owned, not hidden)."],
               ["GET",    "/api/v1/meetings/{id}",                         "owner", "Fetch one meeting."],
               ["PATCH",  "/api/v1/meetings/{id}",                         "owner", "Update title / visibility / recording mode."],
               ["DELETE", "/api/v1/meetings/{id}",                         "owner", "End if active, hide if already closed."],
               ["POST",   "/api/v1/meetings/{id}/reopen",                  "owner", "Reopen a closed meeting (keeps the same room_name)."],
               ["POST",   "/api/v1/meetings/{id}/branding",                "owner", "Upload branding image (multipart)."],
               ["DELETE", "/api/v1/meetings/{id}/branding",                "owner", "Remove branding image."],
               ["POST",   "/api/v1/meetings/{id}/invite",                  "owner", "Send invite emails via Resend."],
               ["POST",   "/api/v1/meetings/{id}/token",                   "owner", "Mint owner LiveKit token + TURN credentials."],
               ["POST",   "/api/v1/meetings/{id}/mute",                    "owner", "Mute one participant."],
               ["POST",   "/api/v1/meetings/{id}/mute-all",                "owner", "Mute every other participant."],
               ["POST",   "/api/v1/meetings/{id}/kick",                    "owner", "Disconnect a participant."],
               ["POST",   "/api/v1/meetings/{id}/presenter",               "owner", "Set or clear room.metadata.presenter_identity."],
               ["POST",   "/api/v1/meetings/{id}/recordings:start",        "owner", "Start LiveKit egress."],
               ["POST",   "/api/v1/meetings/{id}/recordings:stop",         "owner", "Stop LiveKit egress."],
               ["GET",    "/api/v1/discoverable",                          "auth",  "Active meetings (other owners) flagged for at least authenticated discovery."],
               ["GET",    "/api/v1/public-meetings",                       "—",     "Active meetings flagged list_for_anonymous."],
               ["GET",    "/api/v1/rooms/{room_name}/info",                "—",     "Public lobby metadata (title, branding, host name)."],
               ["GET",    "/api/v1/rooms/{room_name}/branding",            "—",     "Public branding image bytes."],
               ["POST",   "/api/v1/rooms/{room_name}/anon-token",          "—",     "Mint anonymous LiveKit token (password-checked)."],
               ["GET/POST/PUT/DELETE", "/api/v1/rooms/{room}/chat[/...]",  "—",     "Chat history, posting, reactions, attachments."],
               ["GET",    "/api/v1/recordings",                            "owner", "List my recordings."],
               ["GET",    "/api/v1/recordings/{id}/download",              "owner", "Stream recording bytes."],
               ["POST",   "/api/v1/recordings/{id}/publish-youtube",       "owner", "Publish to YouTube; deletes local file on success."],
               ["DELETE", "/api/v1/recordings/{id}",                       "owner", "Delete a recording (keeps YouTube link)."],
               ["GET",    "/api/v1/me/preferences",                        "auth",  "Read user prefs (currently: language + manual flag)."],
               ["PUT",    "/api/v1/me/preferences",                        "auth",  "Write user prefs; flips language_set_manually=true on language change."],
               ["POST",   "/api/v1/webhooks/livekit",                      "shared key", "Egress lifecycle (egress_ended → write Recording row)."]],
              col_widths_cm=[2.0, 6.5, 1.5, 6.5])

    add_heading(doc, "3.5 Authentication & SSO", level=2)
    add_para(doc,
        "meet.witysk.org never holds long-lived credentials. Authentication is "
        "fully delegated to one.witysk.org, which issues HS256-signed JWTs "
        "shared via JWT_SECRET_KEY in the meeting-api environment. The SPA "
        "obtains the token by loading one.witysk.org/sso-bootstrap.html in a "
        "hidden iframe and listening for a postMessage. Once received, the "
        "token is mirrored into meet's localStorage so subsequent loads avoid "
        "the iframe round-trip; expiry is detected on 401 and triggers a "
        "single refresh attempt. There is no refresh-token flow in meet "
        "itself — refreshes happen on the one.witysk.org side and are picked "
        "up the next time the iframe runs.")
    add_image(doc, sso_png, caption="Figure 3.3 — SSO bootstrap and authenticated API call.")

    add_heading(doc, "3.6 Internationalization (i18n)", level=2)
    add_para(doc,
        "The SPA ships with one source-of-truth catalog (en.json) of 425 keys "
        "covering 14 namespaces (nav, common, createMeeting, lobby, room, "
        "settings, etc.). 49 additional locales are auto-translated by two "
        "Python utilities under scripts/ that call DeepSeek's chat-completion "
        "API:")
    add_bullets(doc, [
        "translate_locales.py — full-file translation, only writes locales that don't yet exist.",
        "translate_missing_keys.py — delta translator that fills only newly-added keys, preserving existing translations. Used after adding new strings."
    ])
    add_para(doc,
        "Both scripts shield {{interpolation}} tokens and <1>...</1> Trans "
        "wrappers behind <<<KEEP_n>>> sentinels so the model preserves them "
        "exactly. Each locale file is loaded lazily at runtime via "
        "i18next-resources-to-backend, so the bundle never ships unused "
        "languages.")

    add_heading(doc, "3.7 Recording lifecycle", level=2)
    add_image(doc, rec_png, caption="Figure 3.4 — Recording lifecycle.")

    add_heading(doc, "3.8 Deployment topology", level=2)
    add_para(doc,
        "scripts/deploy.sh is a single-command ssh+rsync+compose-up flow. "
        "It rsyncs the working tree to /tmp/meet-stage on the host, syncs "
        "into /opt/meet (preserving .env), then runs docker compose -p meet "
        "up -d --build. It refuses to proceed without an .env at /opt/meet/. "
        "It also explicitly verifies that coturn is still running afterwards "
        "— coturn is in a different process group and our deploys must "
        "never disturb it.")
    add_para(doc,
        "Persistent data lives outside the compose project on the host: "
        "/var/lib/meet/meet.db (SQLite), /var/lib/meet/branding/* (uploaded "
        "images), /var/lib/meet/recordings/* (egress output).",
        italic=True)

    doc.add_page_break()

    # ── 4. Operations ──────────────────────────────────────────────────────
    add_heading(doc, "4. Operations", level=1)

    add_heading(doc, "4.1 Coturn separation", level=2)
    add_para(doc,
        "coturn provides the TURN relay used as a fallback when peers can't "
        "establish direct ICE (typical on restrictive corporate networks). "
        "It is intentionally NOT part of the meet docker compose project. "
        "deploy.sh checks before/after every deploy that coturn is still "
        "active — if it isn't, the script prompts before proceeding. Never "
        "include coturn in this compose project; a frontend-only redeploy "
        "must not have any chance of breaking realtime media for ongoing "
        "calls.")

    add_heading(doc, "4.2 Recording retention", level=2)
    add_para(doc,
        "scheduler.py runs on a periodic loop inside meeting-api. On each "
        "tick it sweeps recordings where expires_at < now() AND status != "
        "'deleted', deletes the local file, and flips status. Default "
        "retention is 30 days; per-user override exists in preferences "
        "(privacy.recordingRetentionDaysOverride) but is not yet honored "
        "server-side.")

    add_heading(doc, "4.3 Backups", level=2)
    add_para(doc,
        "Authoritative state is /var/lib/meet/meet.db (~few MB, SQLite). "
        "Recordings are reconstructable on demand from YouTube once "
        "published, but pre-publish files exist only locally — a "
        "disaster-recovery snapshot of /var/lib/meet covers both.")

    add_heading(doc, "4.4 Monitoring & health", level=2)
    add_para(doc,
        "The deploy script's final step is a curl to /api/health; that "
        "endpoint also serves as a Caddy/Docker healthcheck candidate. "
        "Beyond that there is no formal monitoring; LiveKit exposes its own "
        "Prometheus endpoint that can be scraped.")

    doc.add_page_break()

    # ── 5. Future work ─────────────────────────────────────────────────────
    add_heading(doc, "5. Future work / non-goals", level=1)
    add_bullets(doc, [
        "Move from SQLite to Postgres if the user-prefs / chat tables grow significantly.",
        "Wire the rest of Settings (currently localStorage-only) to the per-user prefs API table.",
        "Switch one.witysk.org to RS256/JWKS so meet doesn't need a shared HS256 secret.",
        "Implement scheduled meetings (scheduled_at exists but is not yet enforced).",
        "Replace the manual 30-day retention sweep with a per-user policy.",
        "Introduce code-splitting in the Vite build (current main bundle is 1.1 MB / 313 KB gzip).",
    ])

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"✓ wrote {OUT_PATH.relative_to(REPO)}  ({OUT_PATH.stat().st_size:,} bytes)")


if __name__ == "__main__":
    build_doc()
